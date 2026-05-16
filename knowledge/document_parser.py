"""
Document Parser - Reads various document formats
Supports: .txt, .pdf, .docx, .xlsx
"""

import os
import logging
from typing import List, Dict
from pathlib import Path

logger = logging.getLogger(__name__)

# Try to import document libraries, handle gracefully if not installed
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


class DocumentParser:
    """Parse various document formats and extract text content."""

    SUPPORTED_EXTENSIONS = [".txt", ".pdf", ".docx", ".xlsx", ""]

    def __init__(self):
        self.supported = {
            ".txt": self._parse_txt,
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".xlsx": self._parse_xlsx,
            "": self._parse_txt,  # Files without extension treated as text
        }

    def parse_file(self, file_path: str) -> str:
        """Parse a single file and return its text content."""
        ext = Path(file_path).suffix.lower()

        if ext not in self.supported:
            logger.warning(f"Unsupported file type: {ext}")
            return ""

        try:
            return self.supported[ext](file_path)
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            return ""

    def parse_directory(self, directory: str = None) -> List[Dict[str, str]]:
        """Parse all supported files in a directory."""
        # Default to knowledge folder in project root
        if directory is None:
            project_root = Path(__file__).parent.parent.parent
            directory = project_root / "knowledge"

        documents = []
        dir_path = Path(directory)

        if not dir_path.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return documents

        for file_path in dir_path.rglob("*"):
            if (
                file_path.is_file()
                and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
            ):
                # Skip README files and Office temporary lock files (~$)
                if ("readme" in file_path.name.lower()
                        or file_path.name.startswith("~$")):
                    continue

                text = self.parse_file(str(file_path))
                if text.strip():
                    documents.append(
                        {
                            "filename": file_path.name,
                            "filepath": str(file_path),
                            "content": text,
                            "extension": file_path.suffix.lower(),
                        }
                    )
                    logger.info(f"Parsed: {file_path.name}")

        return documents

    def _parse_txt(self, file_path: str) -> str:
        """Parse plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file."""
        if PyPDF2 is None:
            logger.warning("PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""

        text = []
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")

        return "\n".join(text)

    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file."""
        if docx is None:
            logger.warning(
                "python-docx not installed. Install with: pip install python-docx"
            )
            return ""

        text = []
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text.append(row_text)
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")

        return "\n".join(text)

    def _parse_xlsx(self, file_path: str) -> str:
        """Parse XLSX file."""
        if openpyxl is None:
            logger.warning("openpyxl not installed. Install with: pip install openpyxl")
            return ""

        text = []
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text.append(f"=== Sheet: {sheet} ===")

                for row in ws.iter_rows(values_only=True):
                    row_values = [str(v) if v is not None else "" for v in row]
                    row_text = " | ".join(row_values)
                    if row_text.strip():
                        text.append(row_text)
        except Exception as e:
            logger.error(f"XLSX parsing error: {e}")

        return "\n".join(text)


# Singleton instance
parser = DocumentParser()
